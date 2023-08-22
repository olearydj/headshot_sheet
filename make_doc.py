import os
import docx
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from rich.traceback import install

install()  # install rich traceback

# Specify your input directory and output docx file here
output_dir = "/Users/djo/Dropbox/Personal/Studies/0-Grad School/MEM Program/IR INSY 3010 - Python SQL/roll-attendance/headshots"
input_dir = output_dir
output_docx = os.path.join(output_dir, "headshots.docx")
placeholder_image = "./placeholder.jpg"
student_list_file = "./all_students.txt"

# Load the student list
with open(student_list_file, "r") as f:
    all_students = [line.strip() for line in f]

print(all_students[0])

# Create a new Word document
doc = docx.Document()

# Create a dictionary mapping names to image paths
image_dict = {}
for filename in os.listdir(input_dir):
    if filename.endswith(".jpg") or filename.endswith(".jpeg"):
        name_parts = filename.split("_")
        first_name = "".join([i for i in name_parts[1] if not i.isdigit()]).title()
        last_name = name_parts[0].title()
        name = f"{first_name} {last_name}"
        image_dict[name] = os.path.join(input_dir, filename)

# Create a 6x6 table for each group of 36 students
for i in range(0, len(all_students), 36):
    table = doc.add_table(rows=6, cols=6)

    # Add an image and a name to each cell
    for j in range(36):
        if i + j < len(all_students):
            # Get the student's name
            name = all_students[i + j]

            # Get the image path or placeholder image if none exists
            image_path = image_dict.get(name, placeholder_image)

            # Add the image to the cell
            cell = table.cell(j // 6, j % 6)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            try:
                run.add_picture(image_path, width=docx.shared.Inches(1))
            except Exception as e:
                print(f"Error adding image: {image_path}")
                print(e)

            # Add the name to the cell
            paragraph = cell.add_paragraph()
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.text = name

    # Add a page break after each table except the last one
    if i + 36 < len(all_students):
        doc.add_page_break()

# Save the document
doc.save(output_docx)
